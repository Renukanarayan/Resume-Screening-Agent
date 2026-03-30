"""Document parser for extracting text from PDF and DOCX files."""

import os
from dataclasses import dataclass
from pathlib import Path
from typing import Literal

# PDF parsing
try:
    from PyPDF2 import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# DOCX parsing
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


@dataclass
class ParseResult:
    """Result of document parsing."""
    text: str
    file_type: Literal["pdf", "docx", "txt", "unknown"]
    page_count: int = 1
    success: bool = True
    error_message: str = ""
    confidence: float = 1.0  # How confident we are in the extraction quality


class DocumentParser:
    """Parse PDF and DOCX documents to extract text content."""
    
    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt"}
    
    def __init__(self):
        """Initialize the document parser."""
        self._check_dependencies()
    
    def _check_dependencies(self) -> None:
        """Check if required parsing libraries are available."""
        if not PDF_AVAILABLE:
            print("Warning: PyPDF2 not installed. PDF parsing will not work.")
        if not DOCX_AVAILABLE:
            print("Warning: python-docx not installed. DOCX parsing will not work.")
    
    def parse(self, file_path: str) -> ParseResult:
        """
        Parse a document and extract its text content.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            ParseResult with extracted text and metadata
        """
        path = Path(file_path)
        
        # Validate file exists
        if not path.exists():
            return ParseResult(
                text="",
                file_type="unknown",
                success=False,
                error_message=f"File not found: {file_path}",
                confidence=0.0
            )
        
        # Get file extension
        ext = path.suffix.lower()
        
        if ext not in self.SUPPORTED_EXTENSIONS:
            return ParseResult(
                text="",
                file_type="unknown",
                success=False,
                error_message=f"Unsupported file type: {ext}. Supported: {self.SUPPORTED_EXTENSIONS}",
                confidence=0.0
            )
        
        # Parse based on file type
        try:
            if ext == ".pdf":
                return self._parse_pdf(path)
            elif ext in {".docx", ".doc"}:
                return self._parse_docx(path)
            elif ext == ".txt":
                return self._parse_txt(path)
            else:
                return ParseResult(
                    text="",
                    file_type="unknown",
                    success=False,
                    error_message=f"No parser available for: {ext}",
                    confidence=0.0
                )
        except Exception as e:
            return ParseResult(
                text="",
                file_type=ext.replace(".", ""),  # type: ignore
                success=False,
                error_message=f"Error parsing file: {str(e)}",
                confidence=0.0
            )
    
    def _parse_pdf(self, path: Path) -> ParseResult:
        """Parse a PDF file."""
        if not PDF_AVAILABLE:
            return ParseResult(
                text="",
                file_type="pdf",
                success=False,
                error_message="PyPDF2 not installed. Run: pip install PyPDF2",
                confidence=0.0
            )
        
        try:
            reader = PdfReader(str(path))
            pages = []
            
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)
            
            full_text = "\n\n".join(pages)
            
            # Estimate confidence based on text quality
            confidence = self._estimate_extraction_confidence(full_text)
            
            return ParseResult(
                text=full_text,
                file_type="pdf",
                page_count=len(reader.pages),
                success=True,
                confidence=confidence
            )
        except Exception as e:
            return ParseResult(
                text="",
                file_type="pdf",
                success=False,
                error_message=f"PDF parsing error: {str(e)}",
                confidence=0.0
            )
    
    def _parse_docx(self, path: Path) -> ParseResult:
        """Parse a DOCX file."""
        if not DOCX_AVAILABLE:
            return ParseResult(
                text="",
                file_type="docx",
                success=False,
                error_message="python-docx not installed. Run: pip install python-docx",
                confidence=0.0
            )
        
        try:
            doc = Document(str(path))
            paragraphs = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        paragraphs.append(row_text)
            
            full_text = "\n".join(paragraphs)
            confidence = self._estimate_extraction_confidence(full_text)
            
            return ParseResult(
                text=full_text,
                file_type="docx",
                page_count=1,  # DOCX doesn't have clear page boundaries
                success=True,
                confidence=confidence
            )
        except Exception as e:
            return ParseResult(
                text="",
                file_type="docx",
                success=False,
                error_message=f"DOCX parsing error: {str(e)}",
                confidence=0.0
            )
    
    def _parse_txt(self, path: Path) -> ParseResult:
        """Parse a plain text file."""
        try:
            # Try different encodings
            for encoding in ["utf-8", "latin-1", "cp1252"]:
                try:
                    text = path.read_text(encoding=encoding)
                    return ParseResult(
                        text=text,
                        file_type="txt",
                        page_count=1,
                        success=True,
                        confidence=1.0  # Plain text is always high confidence
                    )
                except UnicodeDecodeError:
                    continue
            
            return ParseResult(
                text="",
                file_type="txt",
                success=False,
                error_message="Could not decode file with any supported encoding",
                confidence=0.0
            )
        except Exception as e:
            return ParseResult(
                text="",
                file_type="txt",
                success=False,
                error_message=f"Text file reading error: {str(e)}",
                confidence=0.0
            )
    
    def _estimate_extraction_confidence(self, text: str) -> float:
        """
        Estimate how confident we are in the text extraction quality.
        
        Lower confidence when:
        - Text is very short
        - High ratio of special characters (possible OCR issues)
        - Missing expected resume sections
        """
        if not text or len(text) < 100:
            return 0.3
        
        # Check for reasonable text length
        word_count = len(text.split())
        if word_count < 50:
            return 0.5
        
        # Check for garbled text (high special character ratio)
        special_chars = sum(1 for c in text if not c.isalnum() and not c.isspace())
        if len(text) > 0 and special_chars / len(text) > 0.3:
            return 0.6
        
        # Check for common resume keywords
        resume_keywords = ["experience", "education", "skills", "work", "job", "email", "phone"]
        keyword_matches = sum(1 for kw in resume_keywords if kw.lower() in text.lower())
        
        if keyword_matches >= 4:
            return 0.95
        elif keyword_matches >= 2:
            return 0.85
        else:
            return 0.7


# Singleton instance
_parser: DocumentParser | None = None


def get_document_parser() -> DocumentParser:
    """Get the global document parser instance."""
    global _parser
    if _parser is None:
        _parser = DocumentParser()
    return _parser


def parse_document(file_path: str) -> ParseResult:
    """Convenience function to parse a document."""
    return get_document_parser().parse(file_path)
